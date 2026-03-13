"""
File Parser Module
Parses different file types (PDF, DOCX, TXT, Images) and extracts text content
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional
import mimetypes
from io import BytesIO

logger = logging.getLogger(__name__)


def parse_text_file(file_content: bytes, encoding: str = 'utf-8') -> str:
    """
    Parse text file (.txt)
    
    Args:
        file_content: File bytes content
        encoding: Text encoding (default: utf-8)
    
    Returns:
        Extracted text content
    """
    try:
        text = file_content.decode(encoding)
        return text.strip()
    except UnicodeDecodeError:
        # Try alternative encodings
        try:
            text = file_content.decode('latin-1')
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to decode text file: {str(e)}")
            raise ValueError(f"Failed to parse text file: {str(e)}")


def parse_pdf_file(file_content: bytes) -> str:
    """
    Parse PDF file using pdfplumber
    
    Args:
        file_content: PDF file bytes content
    
    Returns:
        Extracted text content
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber is required for PDF parsing. Install with: pip install pdfplumber")
    
    try:
        pdf_file = BytesIO(file_content)
        extracted_text = []
        
        with pdfplumber.open(pdf_file) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                if text:
                    extracted_text.append(f"[Page {page_num}]\n{text}")
                
                # Also extract tables if present
                tables = page.extract_tables()
                if tables:
                    for table_idx, table in enumerate(tables):
                        table_text = "\n".join(["\t".join([str(cell) for cell in row]) for row in table])
                        extracted_text.append(f"[Table {table_idx + 1} on Page {page_num}]\n{table_text}")
        
        if not extracted_text:
            logger.warning("No text extracted from PDF file")
            return "[Empty PDF or unable to extract text]"
        
        return "\n\n".join(extracted_text)
    
    except Exception as e:
        logger.error(f"Error parsing PDF file: {str(e)}")
        raise ValueError(f"Failed to parse PDF file: {str(e)}")


def parse_docx_file(file_content: bytes) -> str:
    """
    Parse DOCX file using python-docx
    
    Args:
        file_content: DOCX file bytes content
    
    Returns:
        Extracted text content
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
    
    try:
        docx_file = BytesIO(file_content)
        document = Document(docx_file)
        
        extracted_text = []
        
        # Extract paragraphs
        for para in document.paragraphs:
            if para.text.strip():
                extracted_text.append(para.text)
        
        # Extract tables
        for table_idx, table in enumerate(document.tables):
            table_text_lines = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_text_lines.append("\t".join(row_cells))
            if table_text_lines:
                extracted_text.append(f"[Table {table_idx + 1}]\n" + "\n".join(table_text_lines))
        
        if not extracted_text:
            logger.warning("No text extracted from DOCX file")
            return "[Empty DOCX or unable to extract text]"
        
        return "\n\n".join(extracted_text)
    
    except Exception as e:
        logger.error(f"Error parsing DOCX file: {str(e)}")
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")


def parse_image_file(file_content: bytes, file_name: str) -> str:
    """
    Parse image file (JPG, PNG, etc.) using OCR (Pillow for image validation)
    For now, returns image file name as context. Can be extended with Tesseract/EasyOCR
    
    Args:
        file_content: Image file bytes content
        file_name: Original file name
    
    Returns:
        Image metadata as text
    """
    try:
        from PIL import Image
    except ImportError:
        logger.warning("Pillow not available for image verification")
    
    try:
        # Verify image file
        image_file = BytesIO(file_content)
        try:
            img = Image.open(image_file)
            width, height = img.size
            format_str = img.format or "unknown"
            return f"[Image File: {file_name}]\nFormat: {format_str}\nDimensions: {width}x{height}\nNote: For better text extraction from images, consider uploading PDF or DOCX with embedded images."
        except Exception:
            # If not a valid image, still process
            return f"[Image/Binary File: {file_name}]\nNote: Could not extract content from this file type. Please ensure it's a valid image file."
    
    except Exception as e:
        logger.error(f"Error processing image file: {str(e)}")
        raise ValueError(f"Failed to process image file: {str(e)}")


def get_file_type(file_name: str) -> str:
    """
    Determine file type from filename extension
    
    Args:
        file_name: Original file name
    
    Returns:
        File type category ('pdf', 'docx', 'text', 'image', 'unknown')
    """
    file_name_lower = file_name.lower()
    
    if file_name_lower.endswith('.pdf'):
        return 'pdf'
    elif file_name_lower.endswith('.docx'):
        return 'docx'
    elif file_name_lower.endswith('.doc'):
        return 'doc'  # Older Word format (not supported yet)
    elif file_name_lower.endswith(('.txt', '.text')):
        return 'text'
    elif file_name_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff')):
        return 'image'
    else:
        # Try MIME type detection
        mime_type, _ = mimetypes.guess_type(file_name)
        if mime_type:
            if 'pdf' in mime_type:
                return 'pdf'
            elif 'word' in mime_type or 'officedocument' in mime_type:
                return 'docx'
            elif 'text' in mime_type:
                return 'text'
            elif 'image' in mime_type:
                return 'image'
        return 'unknown'


def parse_file(file_content: bytes, file_name: str, max_size: int = 10 * 1024 * 1024) -> Dict[str, Any]:
    """
    Main file parser function that dispatches to appropriate parser
    
    Args:
        file_content: File bytes content
        file_name: Original file name
        max_size: Maximum file size allowed (default: 10MB)
    
    Returns:
        {
            'status': 'success' or 'error',
            'file_name': str,
            'file_type': str,
            'content': str (extracted text),
            'message': str (error message if any),
            'char_count': int
        }
    """
    try:
        # Validate file size
        if len(file_content) > max_size:
            return {
                'status': 'error',
                'file_name': file_name,
                'file_type': 'unknown',
                'content': '',
                'message': f'File size exceeds maximum limit of {max_size / (1024*1024):.1f}MB',
                'char_count': 0
            }
        
        # Detect file type
        file_type = get_file_type(file_name)
        
        # Parse based on file type
        extracted_content = None
        
        if file_type == 'pdf':
            extracted_content = parse_pdf_file(file_content)
        
        elif file_type == 'docx':
            extracted_content = parse_docx_file(file_content)
        
        elif file_type == 'text':
            extracted_content = parse_text_file(file_content)
        
        elif file_type == 'image':
            extracted_content = parse_image_file(file_content, file_name)
        
        elif file_type == 'doc':
            # DOC format not currently supported
            extracted_content = f"[Unsupported File Type: {file_name}]\nNote: Only .docx (modern Word format) is supported, not .doc (older format). Please convert your file to .docx."
        
        else:
            # Try to treat as text
            try:
                extracted_content = parse_text_file(file_content)
            except:
                extracted_content = f"[Unknown File Type: {file_name}]\nNote: Unable to parse this file type. Supported formats: PDF, DOCX, TXT, JPG, PNG, GIF, BMP"
        
        return {
            'status': 'success',
            'file_name': file_name,
            'file_type': file_type,
            'content': extracted_content,
            'message': 'File parsed successfully',
            'char_count': len(extracted_content)
        }
    
    except Exception as e:
        logger.error(f"Error in parse_file for {file_name}: {str(e)}")
        return {
            'status': 'error',
            'file_name': file_name,
            'file_type': get_file_type(file_name),
            'content': '',
            'message': f'Failed to parse file: {str(e)}',
            'char_count': 0
        }


def validate_file(file_content: bytes, file_name: str) -> Dict[str, Any]:
    """
    Validate file without parsing full content
    
    Args:
        file_content: File bytes content
        file_name: Original file name
    
    Returns:
        Validation result with file info
    """
    file_type = get_file_type(file_name)
    file_size = len(file_content)
    
    supported_types = ['pdf', 'docx', 'text', 'image']
    is_supported = file_type in supported_types
    
    return {
        'file_name': file_name,
        'file_type': file_type,
        'file_size': file_size,
        'file_size_mb': round(file_size / (1024 * 1024), 2),
        'is_supported': is_supported,
        'max_size_mb': 10
    }
